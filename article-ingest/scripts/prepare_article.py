"""Prepare an article or document for ingestion into a notes vault.

Does the mechanical half and nothing else: resolve the source, extract clean
text with citable position markers, and render the pages whose meaning lives in
figures rather than words. Judgment -- what deserves a note, how to phrase the
claim -- stays with Claude, which reads this script's output.

This script never touches the vault. Every artifact lands in the work
directory; SKILL.md governs what gets written into the vault and where.

Usage:
    python prepare_article.py <url-or-path> [options]

Writes a work directory containing meta.json, article.md, pages/, and
pages.json, then prints a JSON summary to stdout.
"""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
import tempfile
from pathlib import Path

# A Windows console is cp1252 and Python gives stdout a *strict* error handler
# (unlike stderr, which defaults to backslashreplace). Article titles are full
# of em dashes, curly quotes and non-Latin names, so without this the summary
# print dies after all the real work has already succeeded.
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="backslashreplace")
except (AttributeError, ValueError):
    pass

DOC_SUFFIXES = {".pdf", ".docx", ".epub", ".html", ".htm", ".xhtml", ".md",
                ".markdown", ".txt", ".text"}

# Below this, an extraction is almost certainly a paywall stub or a JS shell
# rather than a real article.
THIN_EXTRACTION_WORDS = 150

PAYWALL_SIGNALS = (
    "subscribe to continue", "subscribers only", "create a free account",
    "sign in to read", "this article is for subscribers", "become a member",
    "you have reached your limit", "register to continue", "enable javascript",
)


def die(msg: str, code: int = 1) -> None:
    print(f"ERROR: {msg}", file=sys.stderr)
    sys.exit(code)


def die_missing(package: str, why: str = "") -> None:
    """Exit with an install command naming the interpreter that is running.

    `sys.executable` beats a documented path: whichever python invoked this
    script is the one that needs the package, which is the whole of the fix in
    the common failure -- a venv exists but the script was run with the system
    python, or vice versa.
    """
    detail = f" {why}" if why else ""
    die(f"{package} is not installed.{detail}\n"
        f"Install it with:\n"
        f"  {sys.executable} -m pip install {package}")


def note(msg: str) -> None:
    print(f"  {msg}", file=sys.stderr)


def slugify(text: str, maxlen: int = 60) -> str:
    text = re.sub(r"[^\w\s-]", "", text, flags=re.UNICODE).strip().lower()
    text = re.sub(r"[\s_-]+", "-", text)
    return text[:maxlen].strip("-") or "article"


def word_count(text: str) -> int:
    return len(text.split())


# Producers stamp these into the title/author fields when the real value is
# absent. Taken literally they are worse than nothing, because they suppress
# the fallbacks that would have found a usable name.
JUNK_META = {"", "-", "untitled", "unknown", "none", "n/a",
             "anonymous", "(anonymous)", "author", "title"}


def clean_meta(value: str | None) -> str | None:
    text = (value or "").strip()
    if text.lower() in JUNK_META:
        return None
    # "Microsoft Word - report_final_v3.docx" is a filename, not a title.
    text = re.sub(r"^Microsoft\s+(Word|PowerPoint|Excel)\s*-\s*", "", text)
    text = re.sub(r"\.(docx?|pptx?|xlsx?|pdf|tex)$", "", text, flags=re.I)
    return text.strip() or None


def parse_pdf_date(raw: str | None) -> str | None:
    """PDF dates look like D:20260712165245+00'00'; keep just the day."""
    m = re.match(r"D?:?\s*(\d{4})(\d{2})(\d{2})", (raw or "").strip())
    if not m:
        return None
    year, month, day = m.groups()
    if not ("1900" <= year <= "2200" and "01" <= month <= "12"):
        return None
    return f"{year}-{month}-{day}"


def title_from_layout(page) -> str | None:
    """Recover a title from page 1 by font size.

    A document marks its title typographically, not positionally: the first
    line may be a license stamp, a journal header, or a preprint banner, but
    the title is almost always the largest type on the page. Reading the layout
    beats reading the first line.
    """
    try:
        layout = page.get_text("dict")
    except Exception:
        return None

    spans: list[tuple[float, float, str]] = []
    for block in layout.get("blocks", []):
        for line in block.get("lines", []):
            # Skip rotated text. Preprint servers stamp an identifier down the
            # left margin in large type -- vertical, and not the title.
            direction = line.get("dir") or (1.0, 0.0)
            if abs(float(direction[1])) > 0.1:
                continue
            for span in line.get("spans", []):
                text = (span.get("text") or "").strip()
                if text:
                    top = (span.get("bbox") or [0, 0, 0, 0])[1]
                    spans.append((round(float(span.get("size", 0)), 1), top, text))
    if not spans:
        return None

    largest = max(size for size, _, _ in spans)
    # Allow a little slack so a title split across spans of near-identical size
    # is not truncated to its first fragment.
    picked = [(top, text) for size, top, text in spans if size >= largest - 0.6]
    picked.sort(key=lambda p: p[0])

    title = re.sub(r"\s+", " ", " ".join(text for _, text in picked)).strip()
    return title if 6 <= len(title) <= 200 else None


def title_from_text(text: str) -> str | None:
    """Best-effort title from a document's opening lines.

    Papers downloaded as 2401.12345v2.pdf have a useless filename but a real
    title on line one, so this beats falling straight back to the stem.
    """
    for line in text.strip().splitlines():
        line = line.strip()
        if 10 <= len(line) <= 140 and not line.lower().startswith(("abstract", "http")):
            return line
    return None


def looks_paywalled(text: str) -> bool:
    lowered = text.lower()
    return any(signal in lowered for signal in PAYWALL_SIGNALS)


# --------------------------------------------------------------------------
# Web
# --------------------------------------------------------------------------


# Default library user-agents are widely blocked; presenting as a browser is
# the difference between an article and a 403 on a lot of publisher sites.
BROWSER_UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
              "(KHTML, like Gecko) Chrome/126.0 Safari/537.36")


def fetch_url_bytes(url: str) -> tuple[bytes, str]:
    """Fetch a URL, returning (body, content-type).

    Fetching here rather than inside trafilatura buys the content type, which
    is what distinguishes a link to an article from a link straight at a PDF --
    a very common way to be handed a paper.
    """
    try:
        import requests
    except ImportError:
        die_missing("requests")

    try:
        resp = requests.get(url, headers={"User-Agent": BROWSER_UA},
                            timeout=45, allow_redirects=True)
        resp.raise_for_status()
    except Exception as exc:
        die(f"Could not fetch {url}\n{exc}\n"
            "The host may be blocking automated requests. If you have "
            "legitimate access, use the Chrome fallback described in SKILL.md.")
    return resp.content, resp.headers.get("Content-Type", "").lower()


def url_is_pdf(url: str, content_type: str) -> bool:
    if "application/pdf" in content_type:
        return True
    return url.lower().split("?")[0].split("#")[0].endswith(".pdf")


def extract_html_string(raw: str) -> tuple[str, dict]:
    """Extract an article from HTML source, returning (markdown, metadata)."""
    try:
        import trafilatura
        from trafilatura.settings import use_config
    except ImportError:
        die_missing("trafilatura", "It does the web article extraction.")

    downloaded = raw
    cfg = use_config()
    # Keep the extraction generous: a missed paragraph is worse than a stray nav
    # line, because Claude can ignore boilerplate but cannot recover lost text.
    cfg.set("DEFAULT", "EXTRACTION_TIMEOUT", "0")

    text = trafilatura.extract(
        downloaded, output_format="markdown", include_links=True,
        include_tables=True, include_formatting=True,
        with_metadata=False, favor_recall=True, config=cfg,
    ) or ""

    meta = {}
    try:
        md = trafilatura.extract_metadata(downloaded)
        if md:
            meta = {
                "title": md.title,
                "author": md.author,
                "published": md.date,
                "sitename": md.sitename,
                "description": md.description,
            }
    except Exception:  # metadata is a nicety; never fail the ingest for it
        pass

    return text.strip(), {k: v for k, v in meta.items() if v}


def extract_html_file(path: Path) -> tuple[str, dict]:
    try:
        import trafilatura
    except ImportError:
        die_missing("trafilatura", "It is needed for saved HTML.")

    raw = path.read_text(encoding="utf-8", errors="replace")
    text = trafilatura.extract(
        raw, output_format="markdown", include_links=True,
        include_tables=True, include_formatting=True, favor_recall=True,
    ) or ""
    meta = {}
    try:
        md = trafilatura.extract_metadata(raw)
        if md and md.title:
            meta["title"] = md.title
        if md and md.author:
            meta["author"] = md.author
    except Exception:
        pass
    return text.strip(), meta


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------


def _import_fitz():
    try:
        import pymupdf
        return pymupdf
    except ImportError:
        pass
    try:
        import fitz
        return fitz
    except ImportError:
        die_missing("pymupdf", "It reads and renders PDFs.")


def page_is_visual(page, text: str, scan_drawings: bool) -> str | None:
    """Return why a page is worth rendering, or None if text captured it.

    Three things text extraction cannot carry: raster figures, pages that are
    images of text (scans), and charts drawn as vector strokes -- which extract
    as a scatter of stray axis labels and nothing else.
    """
    if len(text.strip()) < 200:
        # Either a scan or a full-bleed figure; either way, look at it.
        return "little or no text layer"
    try:
        if page.get_images(full=True):
            return "embedded image"
    except Exception:
        pass
    if scan_drawings:
        try:
            # Vector charts produce hundreds of path ops; body text produces
            # almost none, since glyphs are not drawings.
            if len(page.get_drawings()) > 40:
                return "vector graphics (chart or diagram)"
        except Exception:
            pass
    return None


def extract_pdf(path: Path, work: Path, max_render: int,
                want_vision: bool) -> tuple[str, dict, list[dict]]:
    fitz = _import_fitz()
    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        die(f"Could not open PDF: {exc}")

    if doc.is_encrypted and not doc.authenticate(""):
        die("This PDF is password-protected. Decrypt it first, then retry.")

    raw_meta = doc.metadata or {}
    meta = {
        "title": clean_meta(raw_meta.get("title")),  # resolved after text extraction
        "author": clean_meta(raw_meta.get("author")),
        "published": parse_pdf_date(raw_meta.get("creationDate")),
        "pages": doc.page_count,
    }

    # get_drawings() is O(content stream); skip it on long documents where the
    # per-page cost would dominate.
    scan_drawings = doc.page_count <= 80

    parts: list[str] = []
    candidates: list[tuple[int, str]] = []
    total_chars = 0

    for i, page in enumerate(doc, start=1):
        text = page.get_text("text") or ""
        total_chars += len(text.strip())
        # A page marker per page is what makes "p. 7" citable later, the same
        # way timestamps make a video claim checkable.
        parts.append(f"\n\n<!-- page {i} -->\n\n{text.strip()}")
        if want_vision:
            why = page_is_visual(page, text, scan_drawings)
            if why:
                candidates.append((i, why))

    body = "".join(parts).strip()
    meta["has_text_layer"] = total_chars > 50 * max(1, doc.page_count // 10)

    # Metadata titles are unreliable, so work down: typography, then the first
    # usable line, then the filename.
    if not meta["title"]:
        first_page = re.sub(r"<!-- page \d+ -->", "", body[:2000])
        meta["title"] = (
            title_from_layout(doc[0])
            or title_from_text(first_page)
            or path.stem
        )

    pages: list[dict] = []
    if want_vision and candidates:
        out_dir = work / "pages"
        out_dir.mkdir(parents=True, exist_ok=True)

        # Prefer even coverage over the first N pages, so a figure late in the
        # document is not lost to a cap that filled up in the introduction.
        if len(candidates) > max_render:
            step = len(candidates) / max_render
            candidates = [candidates[int(i * step)] for i in range(max_render)]

        for page_no, why in candidates:
            target = out_dir / f"page-{page_no:03d}.png"
            try:
                # 150 DPI keeps axis labels and small-print tables legible
                # without producing images too large to read comfortably.
                pix = doc[page_no - 1].get_pixmap(dpi=150)
                pix.save(str(target))
            except Exception as exc:
                note(f"could not render page {page_no}: {exc}")
                continue
            pages.append({
                "file": target.name,
                "path": str(target),
                "page": page_no,
                "why": why,
            })

    doc.close()
    return body, meta, pages


# --------------------------------------------------------------------------
# DOCX / EPUB / plain
# --------------------------------------------------------------------------


def extract_docx(path: Path) -> tuple[str, dict]:
    try:
        import docx
    except ImportError:
        die_missing("python-docx", "It reads .docx files.")

    doc = docx.Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Preserve heading levels -- they are the document's own structure and
        # the most useful anchor for citing a claim.
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            depth = "".join(c for c in style if c.isdigit()) or "1"
            lines.append(f"\n{'#' * min(int(depth), 6)} {text}\n")
        elif style.startswith("title"):
            lines.append(f"\n# {text}\n")
        else:
            lines.append(text)

    for table in doc.tables:
        rows = ["| " + " | ".join(c.text.strip() for c in r.cells) + " |"
                for r in table.rows]
        if rows:
            lines.append("\n" + "\n".join(rows) + "\n")

    props = doc.core_properties
    meta = {
        "title": (props.title or "").strip() or path.stem,
        "author": (props.author or "").strip() or None,
        "published": props.created.strftime("%Y-%m-%d") if props.created else None,
    }
    return "\n\n".join(lines).strip(), meta


def extract_epub(path: Path) -> tuple[str, dict]:
    try:
        import ebooklib
        from ebooklib import epub
    except ImportError:
        die_missing("ebooklib", "It reads .epub files.")
    try:
        import trafilatura
    except ImportError:
        trafilatura = None

    book = epub.read_epub(str(path))
    chapters: list[str] = []
    for item in book.get_items():
        if item.get_type() != ebooklib.ITEM_DOCUMENT:
            continue
        raw = item.get_content().decode("utf-8", errors="replace")
        if trafilatura:
            text = trafilatura.extract(raw, output_format="markdown",
                                       include_formatting=True) or ""
        else:
            text = html.unescape(re.sub(r"<[^>]+>", " ", raw))
            text = re.sub(r"\s+", " ", text).strip()
        if text.strip():
            # Chapter boundaries are an EPUB's citable positions.
            chapters.append(f"\n\n<!-- chapter: {item.get_name()} -->\n\n{text.strip()}")

    def first(field):
        vals = book.get_metadata("DC", field)
        return vals[0][0] if vals else None

    meta = {
        "title": first("title") or path.stem,
        "author": first("creator"),
        "published": first("date"),
        "chapters": len(chapters),
    }
    return "".join(chapters).strip(), meta


def extract_plain(path: Path) -> tuple[str, dict]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    title = path.stem
    m = re.search(r"^#\s+(.+)$", text, re.M)
    if m:
        title = m.group(1).strip()
    return text, {"title": title}


# --------------------------------------------------------------------------


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("source", help="URL, or path to a pdf/docx/epub/html/md/txt file")
    ap.add_argument("--work-dir", help="Where to stage output (default: a temp dir)")
    ap.add_argument("--max-render", type=int, default=12,
                    help="Cap on PDF pages rendered for the vision pass")
    ap.add_argument("--no-vision", action="store_true",
                    help="Skip page rendering entirely (text only)")
    args = ap.parse_args()

    src_path = Path(args.source)
    is_url = bool(re.match(r"https?://", args.source, re.I))
    is_local = src_path.exists()

    if not is_url and not is_local:
        looks_like_path = (
            src_path.suffix.lower() in DOC_SUFFIXES
            or any(sep in args.source for sep in ("/", "\\"))
        )
        if looks_like_path:
            die(f"No such file: {src_path}\n"
                f"Resolved against the current directory: {Path.cwd()}\n"
                f"Pass an absolute path, or cd to the file's folder first.")
        die(f"Not a URL and not an existing file: {args.source}")

    work = Path(args.work_dir) if args.work_dir else Path(
        tempfile.mkdtemp(prefix="articleingest-")
    )
    work.mkdir(parents=True, exist_ok=True)

    pages: list[dict] = []
    meta: dict = {}

    if is_url:
        note(f"fetching {args.source}")
        raw, content_type = fetch_url_bytes(args.source)
        if url_is_pdf(args.source, content_type):
            note("content type is PDF -- processing as a document")
            downloaded_pdf = work / "download.pdf"
            downloaded_pdf.write_bytes(raw)
            body, meta, pages = extract_pdf(
                downloaded_pdf, work, args.max_render, not args.no_vision
            )
            meta["source_type"] = "pdf (web)"
        else:
            body, meta = extract_html_string(
                raw.decode("utf-8", errors="replace")
            )
            meta["source_type"] = "web"
        meta["url"] = args.source
    else:
        suffix = src_path.suffix.lower()
        note(f"reading {src_path.name} ({suffix or 'no suffix'})")
        if suffix == ".pdf":
            body, meta, pages = extract_pdf(
                src_path, work, args.max_render, not args.no_vision
            )
            meta["source_type"] = "pdf"
        elif suffix == ".docx":
            body, meta = extract_docx(src_path)
            meta["source_type"] = "docx"
        elif suffix == ".epub":
            body, meta = extract_epub(src_path)
            meta["source_type"] = "epub"
        elif suffix in {".html", ".htm", ".xhtml"}:
            body, meta = extract_html_file(src_path)
            meta["source_type"] = "html"
        elif suffix in {".md", ".markdown", ".txt", ".text"}:
            body, meta = extract_plain(src_path)
            meta["source_type"] = "text"
        else:
            die(f"Unsupported file type: {suffix or '(none)'}\n"
                f"Supported: {', '.join(sorted(DOC_SUFFIXES))}")
        meta["source_path"] = str(src_path.resolve())
        meta["url"] = None

    body = (body or "").strip()
    words = word_count(body)

    # Distinguish "nothing came back" from "a stub came back", because the
    # remedies differ: one is a broken source, the other needs the browser.
    thin = words < THIN_EXTRACTION_WORDS
    paywalled = looks_paywalled(body)
    needs_browser = bool(is_url and (thin or paywalled))

    if not body:
        if is_url:
            die("Extraction returned nothing at all. The page is likely rendered "
                "entirely by JavaScript.\nUse the Chrome fallback in SKILL.md, or "
                "save the page and pass the file.")
        die("Extraction returned nothing. The file may be empty or corrupt.")

    if meta.get("source_type") == "pdf" and not meta.get("has_text_layer", True):
        note("this PDF has no usable text layer -- it is a scan; "
             "rendered pages carry the content")

    meta.setdefault("title", src_path.stem if is_local else args.source)
    meta["slug"] = slugify(meta.get("title") or "article")
    meta["word_count"] = words
    meta["source"] = args.source

    header = f"# {meta['title']}\n"
    if meta.get("author"):
        header += f"\n*{meta['author']}*"
    if meta.get("url"):
        header += f"\n\nSource: {meta['url']}"
    (work / "article.md").write_text(f"{header}\n\n{body}\n", encoding="utf-8")

    (work / "pages.json").write_text(
        json.dumps(pages, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    (work / "meta.json").write_text(
        json.dumps(meta, indent=2, ensure_ascii=False), encoding="utf-8"
    )

    summary = {
        "work_dir": str(work),
        "title": meta.get("title"),
        "slug": meta["slug"],
        "author": meta.get("author"),
        "published": meta.get("published"),
        "source_type": meta["source_type"],
        "url": meta.get("url"),
        "word_count": words,
        "pages": meta.get("pages"),
        "has_text_layer": meta.get("has_text_layer"),
        "rendered_pages": len(pages),
        "thin_extraction": thin,
        "paywall_suspected": paywalled,
        "needs_browser": needs_browser,
        "article_md": str(work / "article.md"),
        "pages_json": str(work / "pages.json"),
    }
    # ensure_ascii keeps the summary printable under any console encoding while
    # staying lossless; meta.json on disk holds the characters verbatim.
    print(json.dumps(summary, indent=2, ensure_ascii=True))


if __name__ == "__main__":
    main()
